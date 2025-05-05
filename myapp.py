import streamlit as st
from docx import Document
import os
import re
import sys
from datetime import datetime

# ======================
# UTILITY FUNCTIONS
# ======================

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def replace_text_in_document(template_path, output_path, replacements):
    """ Replace placeholders in Word document """
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace(key, value)

    doc.save(output_path)

def get_first_word(text):
    words = re.split(r'[\s:,-]', text)
    for word in words:
        if word.strip():
            return word.strip()
    return "Document"

# ======================
# MAIN APP
# ======================

def main():
    st.title("📄 Lab Report Cover Page Generator")
    st.markdown("---")

    with st.form("user_form"):
        name = st.text_input("👤 Name:")
        reg_no = st.text_input("🆔 Registration Number:")
        semester = st.text_input("🎓 Semester:")
        num_templates = st.number_input("📄 Number of Documents to Generate:", min_value=1, step=1, value=1)

        titles = []
        dates = []
        templates_selected = []

        for i in range(num_templates):
            st.markdown(f"### 📄 Document {i+1}")
            title = st.text_input(f"Title for Document {i+1}:", key=f"title_{i}")
            format_choice = st.radio(
                f"Title format for Document {i+1}:",
                ["Single line", "Two lines"],
                key=f"format_{i}"
            )
            date = st.text_input(
                f"Date for Document {i+1} (DD/MM/YYYY):",
                value="01/01/2024",
                key=f"date_{i}"
            )

            titles.append(title)
            dates.append(date)
            templates_selected.append("template1.docx" if format_choice == "Two lines" else "template2.docx")

        submitted = st.form_submit_button("🚀 Generate Documents")

    if submitted:
        if not all([name, reg_no, semester]):
            st.error("❗ Please fill in all user fields")
            return

        output_dir = os.path.abspath(".")
        created_files = []

        for i in range(num_templates):
            title = titles[i]
            date = dates[i]
            template_path = resource_path(templates_selected[i])

            if not title or not date:
                st.warning(f"⚠ Document {i+1} skipped due to missing title or date.")
                continue

            if not os.path.exists(template_path):
                st.error(f"⛔ Template file not found: {templates_selected[i]}")
                continue

            replacements = {
                "{{TITLE}}": title,
                "{{NAME}}": name,
                "{{REG_NO}}": reg_no,
                "{{SEMESTER}}": semester,
                "{{DATE}}": date
            }

            first_word = get_first_word(title)
            base_filename = f"{first_word}.docx"
            output_filename = base_filename
            counter = 1

            while os.path.exists(os.path.join(output_dir, output_filename)):
                output_filename = f"{first_word}_{counter}.docx"
                counter += 1

            output_path = os.path.join(output_dir, output_filename)

            try:
                replace_text_in_document(template_path, output_path, replacements)
                if os.path.exists(output_path):
                    created_files.append(output_filename)
                    st.success(f"✅ Created: {output_filename}")
                    with open(output_path, "rb") as file:
                        st.download_button(
                            label=f"⬇ Download {output_filename}",
                            data=file,
                            file_name=output_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_{i}"
                        )
                else:
                    st.error(f"❌ File not created: {output_path}")
            except Exception as e:
                st.error(f"❌ Failed to create document {i+1}: {str(e)}")

        if created_files:
            st.balloons()
            st.success(f"🎉 Successfully created {len(created_files)} documents!")
        else:
            st.info("📭 No documents were created.")

if __name__ == "__main__":
    main()
